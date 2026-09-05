"""Reading a Word document: what survives the crossing into HTML.

Every document here is a real one written by python-docx, because what is being
tested is what comes back out of the format.

The ways a document can fail to be read at all live beside this in
``test_word_failures.py``.
"""

from __future__ import annotations

from pathlib import Path

import docx

from plainsight.infrastructure.word_reader import WordDocumentReader


def a_word_file(tmp_path: Path, build) -> str:
    """A real .docx on disk, built by whatever the test wants in it."""
    document = docx.Document()
    build(document)
    path = tmp_path / "report.docx"
    document.save(str(path))
    return str(path)


def test_a_heading_crosses_as_a_heading(tmp_path: Path) -> None:
    path = a_word_file(tmp_path, lambda d: d.add_heading("The Title", level=1))

    body = WordDocumentReader().read_body(path)

    assert "<h1>The Title</h1>" in body.text
    assert not body.failure


def test_a_deeper_heading_keeps_its_depth(tmp_path: Path) -> None:
    """Taken from the digit in the style name, so it survives a translation."""

    def build(document) -> None:
        document.add_heading("Top", level=1)
        document.add_heading("Under it", level=3)

    body = WordDocumentReader().read_body(a_word_file(tmp_path, build))

    assert "<h1>Top</h1>" in body.text
    assert "<h3>Under it</h3>" in body.text


def test_paragraphs_come_across_in_order(tmp_path: Path) -> None:
    def build(document) -> None:
        document.add_paragraph("First paragraph.")
        document.add_paragraph("Second paragraph.")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert text.index("First paragraph.") < text.index("Second paragraph.")


def test_a_bulleted_list_crosses_as_one_list_of_items(tmp_path: Path) -> None:
    """One list element around the run of them, not one around each."""

    def build(document) -> None:
        document.add_paragraph("Milk", style="List Bullet")
        document.add_paragraph("Eggs", style="List Bullet")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<ul>" in text
    assert text.count("<ul>") == 1
    assert "<li>Milk</li>" in text
    assert "<li>Eggs</li>" in text
    assert text.index("<ul>") < text.index("<li>Milk</li>")
    assert text.index("<li>Eggs</li>") < text.index("</ul>")


def test_a_numbered_list_crosses_as_a_numbered_list(tmp_path: Path) -> None:
    def build(document) -> None:
        document.add_paragraph("Step one", style="List Number")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<ol>" in text
    assert "<li>Step one</li>" in text


def test_a_paragraph_after_a_list_closes_the_list(tmp_path: Path) -> None:
    """An item and the paragraph under it are not the same thing."""

    def build(document) -> None:
        document.add_paragraph("Milk", style="List Bullet")
        document.add_paragraph("Then the shopping was done.")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert text.index("</ul>") < text.index("Then the shopping was done.")


def test_bold_and_italic_survive(tmp_path: Path) -> None:
    def build(document) -> None:
        paragraph = document.add_paragraph()
        paragraph.add_run("plain ")
        paragraph.add_run("strong").bold = True
        paragraph.add_run(" and ")
        paragraph.add_run("leaning").italic = True

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<strong>strong</strong>" in text
    assert "<em>leaning</em>" in text


def test_one_phrase_split_across_runs_is_marked_once(tmp_path: Path) -> None:
    """Word splits a run wherever it likes; the emphasis is still one phrase.

    Marked run by run it became three elements where the author wrote one
    word. Measured on a real document before this: three such pairs.
    """

    def build(document) -> None:
        paragraph = document.add_paragraph()
        for piece in ("imp", "ort", "ant"):
            paragraph.add_run(piece).bold = True

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<strong>important</strong>" in text
    assert text.count("<strong>") == 1


def test_a_space_inside_a_run_is_not_given_emphasis_of_its_own(
    tmp_path: Path,
) -> None:
    """Emphasis on a space shows as nothing and says nothing."""

    def build(document) -> None:
        paragraph = document.add_paragraph()
        paragraph.add_run("before")
        paragraph.add_run(" bold ").bold = True
        paragraph.add_run("after")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "before<strong> bold </strong>after" in text


def test_a_character_that_is_markup_is_escaped_rather_than_obeyed(
    tmp_path: Path,
) -> None:
    """This is the whole reason the reader hands over HTML rather than Markdown.

    Escaping into HTML is total: five characters mean something and each has a
    named escape, so text cannot be read as markup whatever the author typed.
    Markdown has no such guarantee, which is how an indent became a block of
    code.
    """

    def build(document) -> None:
        document.add_paragraph('Use <b>tags</b> & ampersands & quotes " freely')

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "&lt;b&gt;tags&lt;/b&gt;" in text
    assert "&amp; ampersands" in text
    assert "<b>" not in text


def test_a_paragraph_that_is_markdown_is_not_obeyed_either(tmp_path: Path) -> None:
    """A hyphen, a hash and an indent are all just characters the author typed."""

    def build(document) -> None:
        document.add_paragraph("# Not a heading")
        document.add_paragraph("- Not an item")
        document.add_paragraph("    Not a block of code")
        document.add_paragraph("A snake_case_name and a *star*")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<p># Not a heading</p>" in text
    assert "<p>- Not an item</p>" in text
    assert "<p>Not a block of code</p>" in text
    assert "<p>A snake_case_name and a *star*</p>" in text
    assert "<em>" not in text


def test_a_layout_table_gives_up_its_contents_rather_than_a_table(
    tmp_path: Path,
) -> None:
    """Word arranges pages with tables as often as it tabulates anything.

    Measured on a real CV: one single-row layout table became a table row 2550
    characters long, with every paragraph of a whole section run together on
    one line, because a cell's own line breaks were collapsed to fit it.
    """

    def build(document) -> None:
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A heading\n\nA paragraph beneath it."
        table.cell(0, 1).text = "Something beside it."

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<table>" not in text
    assert "<p>A heading</p>" in text
    assert "<p>A paragraph beneath it.</p>" in text
    assert "<p>Something beside it.</p>" in text


def test_a_table_of_data_crosses_as_a_table(tmp_path: Path) -> None:
    def build(document) -> None:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Key"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "a"
        table.cell(1, 1).text = "1"

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<table>" in text
    assert "<tr><th>Key</th><th>Value</th></tr>" in text
    assert "<tr><td>a</td><td>1</td></tr>" in text


def test_a_table_lands_between_the_paragraphs_it_sits_between(
    tmp_path: Path,
) -> None:
    """Read off two separate lists, a table lands after every paragraph.

    python-docx exposes paragraphs and tables as two collections, each in
    document order only within itself. Walking the body's own XML is what keeps
    a table where its author put it.
    """

    def build(document) -> None:
        document.add_paragraph("Before the table.")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "In the table"
        document.add_paragraph("After the table.")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert text.index("Before the table.") < text.index("In the table")
    assert text.index("In the table") < text.index("After the table.")


def test_an_empty_paragraph_leaves_nothing_behind_it(tmp_path: Path) -> None:
    def build(document) -> None:
        document.add_paragraph("Something.")
        document.add_paragraph("")
        document.add_paragraph("Something else.")

    text = WordDocumentReader().read_body(a_word_file(tmp_path, build)).text

    assert "<p></p>" not in text
    assert text.count("<p>") == 2


def test_a_wall_of_a_paragraph_is_broken_into_paragraphs_of_its_own(
    tmp_path: Path,
) -> None:
    """The rule is the domain's; what changes is only the shape it comes in.

    On a surface that renders Markdown the passages are run together with a
    gap between them. Here each one is a paragraph, which is the same break in
    the place HTML keeps it.
    """
    sentence = "This is a sentence of a reasonable length that says something. "
    wall = sentence * 12

    path = a_word_file(tmp_path, lambda d: d.add_paragraph(wall))

    text = WordDocumentReader().read_body(path).text

    assert text.count("<p>") > 1
    rejoined = text.replace("</p>\n<p>", "").replace("<p>", "").replace("</p>", "")
    assert rejoined == wall.strip()


def test_a_wall_is_cut_between_characters_never_inside_an_escape(
    tmp_path: Path,
) -> None:
    """A break placed in finished HTML could split `&amp;` into rubbish.

    The passages are found in the text the author wrote, before a character is
    escaped or a tag put near it, so a cut can only fall between two whole
    characters of the original.
    """
    sentence = "Integration & platform work, done at the seams of systems. "
    wall = sentence * 12

    path = a_word_file(tmp_path, lambda d: d.add_paragraph(wall))

    text = WordDocumentReader().read_body(path).text

    assert text.count("<p>") > 1
    assert "&amp;" in text
    assert "&am" not in text.replace("&amp;", "")


def test_a_heading_is_never_broken_however_long_it_runs(tmp_path: Path) -> None:
    """A heading is one thing by definition; broken in two it is two headings."""
    long_heading = "A heading that simply will not stop going on and on. " * 12

    path = a_word_file(tmp_path, lambda d: d.add_heading(long_heading, level=1))

    text = WordDocumentReader().read_body(path).text

    assert text.count("<h1>") == 1
